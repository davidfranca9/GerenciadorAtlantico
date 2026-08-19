"""Geracao de documentos DOCX (Ordem de Coleta / Carta Frete).

Portado de gerenciador_atlantico/servicos/documentos.py, removendo a
dependencia de tkinter/shared.py para poder rodar em um servidor Linux.
"""
from __future__ import annotations

import locale
import re
import unicodedata
from datetime import datetime

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

SHEET_NAME = "Ordem de Carregamento"
AUTORIZACAO_FIRST_DATA_ROW = 3
AUTORIZACAO_TOTAL_ROW = 15
TOTAL_LABEL = "Peso total (t)"

OC_HEADERS = ["Pedido", "Produto", "Embalagem", "Peso (t)", "Cidade/UF", "Cliente"]
OC_COLUMN_WIDTHS_IN = [0.55, 1.55, 0.78, 0.58, 1.05, 1.39]

LABEL_PATTERN = re.compile(
    "((?:Motorista|CNH|Fone|Telefone)|(?:(?:1(?:a|ª)?|2(?:a|ª)?|3(?:a|ª)?)\\s*Placa))",
    re.IGNORECASE,
)
STANDARDIZED_LABELS = {
    "motorista": "Motorista",
    "cnh": "CNH",
    "fone": "Fone",
    "1": "1a Placa",
    "2": "2a Placa",
    "3": "3a Placa",
}


def _clean(s):
    return re.sub(r"\s+", " ", str(s)).strip() if s is not None else ""


def _format_peso(v):
    if v is None:
        return ""
    try:
        v_str = str(v).replace(",", ".")
        f = float(v_str)
        formatted_str = f"{f:.3f}".rstrip("0").rstrip(".")
        return formatted_str or "0"
    except (ValueError, TypeError):
        return _clean(v)


def _format_peso_documento(value):
    text = _format_peso(value)
    if re.fullmatch(r"\d+(?:\.\d+)?", text):
        return text.replace(".", ",")
    return text


def formatar_moeda_brasileira(valor_str: str) -> str:
    if not valor_str:
        return ""
    try:
        try:
            locale.setlocale(locale.LC_ALL, "pt_BR.UTF-8")
        except locale.Error:
            pass
        valor_limpo = valor_str.replace(".", "").replace(",", ".")
        valor_float = float(valor_limpo)
        return f"{valor_float:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except ValueError:
        return valor_str


def _safe_float(value):
    if value is None:
        return 0.0
    text = str(value).strip()
    if not text:
        return 0.0
    text = text.replace(" t", "").replace("T", "").strip()
    if "," in text and "." in text:
        text = text.replace(".", "").replace(",", ".")
    else:
        text = text.replace(",", ".")
    try:
        return float(text)
    except (TypeError, ValueError):
        return 0.0


def calcular_peso_total(produtos):
    return sum(_safe_float(p.get("toneladas")) for p in produtos or [])


def _find_prod_table(doc):
    for t in doc.tables:
        if t.rows and len(t.rows[0].cells) >= 2:
            header = [c.text.strip().lower() for c in t.rows[0].cells]
            if "pedido" in header[0] and "produto" in header[1]:
                return t
    return None


def _label_key_from_text(text):
    text = text.strip().lower()
    if "motorista" in text:
        return "motorista"
    if "cnh" in text:
        return "cnh"
    if "fone" in text or "telefone" in text:
        return "fone"
    if "placa" in text:
        if re.search(r"^\s*1", text):
            return "1"
        if re.search(r"^\s*2", text):
            return "2"
        if re.search(r"^\s*3", text):
            return "3"
    return None


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, value, bold=False, size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(_clean(value)))
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_row_fill(row, fill):
    for cell in row.cells:
        _set_cell_shading(cell, fill)


def _apply_oc_table_geometry(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for row in table.rows:
        for idx, width in enumerate(OC_COLUMN_WIDTHS_IN):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def _replace_paragraph_text(paragraph, replacements):
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    updated = full_text
    for key, value in replacements.items():
        updated = updated.replace(key, str(value or ""))
    if updated == full_text:
        return
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""


def _iter_all_paragraphs(parent):
    for paragraph in getattr(parent, "paragraphs", []):
        yield paragraph
    for table in getattr(parent, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def fill_products_in_existing_table(doc, produtos):
    table = _find_prod_table(doc)
    if not table:
        return

    while len(table.columns) < len(OC_HEADERS):
        table.add_column(Inches(0.45))

    _apply_oc_table_geometry(table)

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    header_row = table.rows[0]
    for cell, header in zip(header_row.cells, OC_HEADERS):
        _set_cell_text(cell, header, bold=True, size=7.5, color="000000")

    for p in produtos or []:
        row = table.add_row()
        values = [
            p.get("contrato", ""),
            p.get("produto", ""),
            p.get("embalagem", ""),
            _format_peso_documento(p.get("toneladas")),
            p.get("cidade", ""),
            p.get("cliente", ""),
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx in (1, 5) else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(row.cells[idx], value, size=7.2, align=align)

    total_row = table.add_row()
    _set_row_fill(total_row, "E6F4F1")
    try:
        label_cell = total_row.cells[0].merge(total_row.cells[2])
    except Exception:
        label_cell = total_row.cells[0]
    _set_cell_text(label_cell, "PESO TOTAL", bold=True, size=7.5, align=WD_ALIGN_PARAGRAPH.RIGHT, color="0F3D36")
    _set_cell_text(total_row.cells[3], _format_peso_documento(calcular_peso_total(produtos)), bold=True, size=7.5, color="0F3D36")
    for idx in (4, 5):
        if idx < len(total_row.cells):
            _set_cell_text(total_row.cells[idx], "", size=7.2)

    _apply_oc_table_geometry(table)


def copy_run_style(src_run, dest_run):
    try:
        dest_run.font.name = src_run.font.name
        dest_run.font.size = src_run.font.size
        dest_run.font.bold = src_run.font.bold
        dest_run.font.italic = src_run.font.italic
        dest_run.font.underline = src_run.font.underline
        if src_run.font.color and src_run.font.color.rgb:
            dest_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass


def fill_motorista_and_placas(doc, cpf, nome, cnh, fone, placa1, placa2, placa3):
    motorista = " - ".join(part for part in [cpf, nome] if str(part or "").strip())
    mapping = {"motorista": motorista, "cnh": cnh, "fone": fone, "1": placa1, "2": placa2, "3": placa3}
    for para in _iter_all_paragraphs(doc):
        if ":" not in para.text:
            continue
        matches = list(LABEL_PATTERN.finditer(para.text))
        if matches:
            src_run = para.runs[0] if para.runs else None
            for run in para.runs:
                run.text = ""
            for idx, m in enumerate(matches):
                key = _label_key_from_text(m.group(0))
                if not key:
                    continue
                if idx > 0:
                    sep_run = para.add_run("\t\t")
                    if src_run:
                        copy_run_style(src_run, sep_run)
                val = mapping.get(key, "")
                label = STANDARDIZED_LABELS.get(key)
                label_run = para.add_run(f"{label}:")
                if src_run:
                    copy_run_style(src_run, label_run)
                label_run.font.bold = True
                if val:
                    val_run = para.add_run(f" {val}")
                    if src_run:
                        copy_run_style(src_run, val_run)
                    val_run.font.bold = False


def _replace_date_in_paragraph(paragraph, new_date):
    for run in paragraph.runs:
        if re.search(r"\d{1,2}/\d{1,2}/\d{2,4}", run.text):
            run.text = re.sub(r"\d{1,2}/\d{1,2}/\d{2,4}", new_date, run.text, 1)
            return True
    return False


def gerar_oc_docx(modelo_path, save_path, produtos, cpf, nome, cnh, fone, placa1, placa2, placa3, data_carregamento):
    doc = Document(modelo_path)
    fill_products_in_existing_table(doc, produtos)
    fill_motorista_and_placas(doc, cpf, nome, cnh, fone, placa1, placa2, placa3)

    hoje = datetime.now().strftime("%d/%m/%Y")
    for p in _iter_all_paragraphs(doc):
        if "Emiss" in p.text and re.search(r"\d{1,2}/\d{1,2}/\d{2,4}", p.text):
            _replace_date_in_paragraph(p, hoje)
            break

    doc.save(save_path)


def fill_carta_frete_docx(doc, dados):
    valor_frete_str = str(dados.get("VALOR_FRETE", ""))
    valor_formatado = formatar_moeda_brasileira(valor_frete_str) if valor_frete_str else ""
    valor_com_moeda = f"R$ {valor_formatado}" if valor_formatado else ""
    replacements = {
        "{{DATA}}": dados.get("DATA", ""),
        "{{CONDUTOR}}": dados.get("CONDUTOR", ""),
        "{{CPF}}": dados.get("CPF", ""),
        "{{PLACA_CAVALO}}": dados.get("PLACA_CAVALO", ""),
        "{{VALOR_FRETE}}": valor_com_moeda,
        "{{AUTORIZACAO_NUM}}": dados.get("AUTORIZACAO_NUM", ""),
    }
    used_placeholders = False
    for paragraph in _iter_all_paragraphs(doc):
        if any(key in paragraph.text for key in replacements):
            used_placeholders = True
            _replace_paragraph_text(paragraph, replacements)
    if used_placeholders:
        return

    if valor_frete_str:
        for table in doc.tables:
            encontrado = False
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "R$" in p.text and valor_formatado not in p.text:
                            run = p.add_run(" " + valor_formatado)
                            run.font.name = "Calibri (Corpo)"
                            run.font.size = Pt(14)
                            run.font.bold = True
                            encontrado = True
                            break
                    if encontrado:
                        break
                if encontrado:
                    break
            if encontrado:
                break

    mapa_campos_normais = {
        "DATA": "DATA:",
        "CONDUTOR": "CONDUTOR:",
        "CPF": "CPF:",
        "PLACA_CAVALO": "PLACA CAVALO:",
        "AUTORIZACAO_NUM": "AUTORIZAÇÃO Nº:",
    }

    def preencher_tabela(table):
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                for subtable in cell.tables:
                    preencher_tabela(subtable)
                for p in cell.paragraphs:
                    for chave, rotulo in mapa_campos_normais.items():
                        valor = str(dados.get(chave, ""))
                        if rotulo in p.text and valor not in p.text:
                            if ci + 1 < len(row.cells):
                                target_cell = row.cells[ci + 1]
                                target_cell.text = ""
                                run = target_cell.add_paragraph(valor).runs[0]
                                run.bold = True
                            else:
                                p.add_run(" " + valor).bold = True

    for table in doc.tables:
        preencher_tabela(table)


def get_headers_from_sheet(ws):
    return [str(h.value) if h.value is not None else "" for h in ws[1]]


def _format_autorizacao_sheet(ws, headers, total_row=AUTORIZACAO_TOTAL_ROW):
    header_fill = PatternFill("solid", fgColor="0F766E")
    total_fill = PatternFill("solid", fgColor="E6F4F1")
    thin = Side(style="thin", color="C7D2D0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col_idx, header in enumerate(headers, start=1):
        for row_idx in (1, 2):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.fill = header_fill
            cell.font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border

        width = {
            "Cliente": 34,
            "Data de Carregamento": 18,
            "Placa cavalo mecânico": 18,
            "Nome do condutor": 28,
            "Número do pedido": 16,
            "Produto": 34,
            "Embalagem": 16,
            "Quantidade": 14,
            "Cidade/UF": 22,
        }.get(header, 16)
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    for row_idx in range(AUTORIZACAO_FIRST_DATA_ROW, total_row + 1):
        ws.row_dimensions[row_idx].height = 22
        for col_idx, _header in enumerate(headers, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border

    try:
        qty_col_idx = headers.index("Quantidade") + 1
    except ValueError:
        return

    label_col_idx = max(1, qty_col_idx - 1)
    label_cell = ws.cell(row=total_row, column=label_col_idx)
    total_cell = ws.cell(row=total_row, column=qty_col_idx)
    label_cell.value = TOTAL_LABEL
    total_cell.value = f"=SUM({get_column_letter(qty_col_idx)}{AUTORIZACAO_FIRST_DATA_ROW}:{get_column_letter(qty_col_idx)}{total_row - 1})"
    total_cell.number_format = "0.###"

    for cell in (label_cell, total_cell):
        cell.fill = total_fill
        cell.font = Font(name="Arial", size=10, bold=True, color="0F3D36")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border


def _clear_autorizacao_data_rows(ws, headers, total_row=AUTORIZACAO_TOTAL_ROW):
    for row_num in range(AUTORIZACAO_FIRST_DATA_ROW, total_row):
        for col_num in range(1, len(headers) + 1):
            ws.cell(row=row_num, column=col_num).value = None


def _write_autorizacao_rows(ws, headers, produtos, data_carregamento, nome_condutor=None, placa_cavalo=None):
    max_items = max(0, AUTORIZACAO_TOTAL_ROW - AUTORIZACAO_FIRST_DATA_ROW)
    for idx, p in enumerate((produtos or [])[:max_items]):
        contrato = str(p.get("contrato") or "").strip()
        row_map = {
            "Cliente": p.get("cliente"),
            "Data de Carregamento": data_carregamento,
            "Placa cavalo mecânico": placa_cavalo,
            "Nome do condutor": nome_condutor,
            "Número do pedido": int(contrato) if contrato.isdigit() else contrato,
            "Produto": p.get("produto"),
            "Embalagem": p.get("embalagem"),
            "Quantidade": _safe_float(p.get("toneladas")),
            "Cidade/UF": p.get("cidade"),
        }
        for col_idx, header in enumerate(headers, start=1):
            ws.cell(row=AUTORIZACAO_FIRST_DATA_ROW + idx, column=col_idx).value = row_map.get(header)


def gerar_autorizacao_xlsx(template_path, save_path, produtos, data_carregamento, nome_condutor, placa_cavalo):
    wb = load_workbook(template_path)
    ws = wb[SHEET_NAME] if SHEET_NAME in wb.sheetnames else wb.active
    headers = get_headers_from_sheet(ws)
    _clear_autorizacao_data_rows(ws, headers)
    _write_autorizacao_rows(ws, headers, produtos, data_carregamento, nome_condutor, placa_cavalo)
    _format_autorizacao_sheet(ws, headers)
    try:
        wb.calculation.fullCalcOnLoad = True
    except Exception:
        pass
    wb.save(save_path)
